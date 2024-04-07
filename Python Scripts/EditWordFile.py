import docx
import os
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import dropbox
import dropbox.files


with open("<TOKEN FILE>","r") as f:
        TOKEN = f.read()
dbx = dropbox.Dropbox(TOKEN)
dbx.check_and_refresh_access_token()

def delete_files():
    """Deletes the text file from the file buffer"""
   
    try:
        folder_path="<FILE BUFFER PATHWAY>"
        # List all files in the specified folder
        files = os.listdir(folder_path)

        # Find and delete files with the specified pattern
        for file in files:
            if file.endswith("createInvoice.txt"):
                file_path = os.path.join(folder_path, file)
                os.remove(file_path)
                print("createInvoice.txt is deleted")

    except Exception as e:
        print(f"An error occurred: {e}")

def upload_files_to_folder(file_name):
    """Uploads all the files from the buffer folder to a specific folder in Dropbox"""
    for file in os.listdir("<EMAIL BUFFER PATHWAY>"):
        if file.endswith(".DS_Store") is False:
            with open(os.path.join("<EMAIL BUFFER PATHWAY>",file), "rb") as f:
                data = f.read()
                dbx.files_upload(data, file_name+f"{file}")


def extract_data_from_text_file(folder_path):
    """Gets Client details from folder to generate invoice"""
    try:
        # List all files in the specified folder
        files = os.listdir(folder_path)
        data = []

        # Find the first text file in the folder
        for file in files:
            if file.endswith("createInvoice.txt"):
                text_file_path = os.path.join(folder_path, file)

                # Read and extract data from the text file
                with open(text_file_path, 'r') as file_content:
                    data = file_content.readlines()
                    for item in data:
                        item.strip()
                    return data

        # If no text file is found
        print("No text file found in the folder.\n\n\n")
        return None

    except Exception as e:
        print(f"An error occurred: {e}")
        return None


def edit_word_file(file_path, data):
    """Generates an invoice to be stored and emailed to client"""
    doc = docx.Document("<INVOICE TEMPLATE PATHWAY>")

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key,value)

    # Making the table data 
    table_header = ["Goods", "Value"]
    amount = "R "+data["[Amount]"]
    session = "Coaching Session "+ data["[Date]"]
    table_data = [
        [session,amount], 
        ["TOTAL DUE",amount]
    ]

    table = doc.add_table(rows=1,cols=2)
    for i in range(2):
        table.rows[0].cells[i].text = table_header[i]
    
    for Goods, Value in table_data:
        cells = table.add_row().cells
        cells[0].text = Goods
        cells[1].text = Value
    

    # Formatting table (Dear god this was waaay more difficult than it should have been)
    for i in range(3):
        for j in range(2):   
            cells_xml_element = table.rows[i].cells[j]._tc
            table_cell_properties = cells_xml_element.get_or_add_tcPr()

            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'),'bfbfbf') 

            table_cell_properties.append(shade_obj)

    # Adding stuff after table 
    doc.add_paragraph("\n\nPAYMENT DUE via YOCO. (Link sent to mobile number)\nOr EFT before session\n\n<BANK DETAILS>")
    doc.add_paragraph("Please send proof of payment to <EMAIL>")
    

    # run = doc.add_paragraph().add_run("")
    # font = run.font
    # font.color.rgb = RGBColor(0,0,255)

    doc.save(file_path)

def main():
        data = extract_data_from_text_file("/Users/CrispyBacon/Desktop/Mumsle App/FileBuffer/")
        file_name=data[0].strip()+" Invoice "+data[2].strip()
        folder = f"/Clients/{data[0].strip()}/"
        data = {
        "[Name]": data[0],"[Date]": data[1],
        "[Number]":data[2],
        "[Amount]":data[3]
        }
        file_name = "<EMAIL BUFFER PATHWAY>"+file_name+".docx"
        edit_word_file(file_name,data)

        delete_files()
        upload_files_to_folder(folder)


if __name__ == "__main__":
    main()


